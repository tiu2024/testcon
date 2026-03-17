#!/usr/bin/env python3
"""
Quiz Processor - Converts multiple choice questions from text to Word documents and new format.

Generates output files based on flags:
    -f  / --table-format:    Word document with questions in table format
    -n  / --hash-format:     Text file with {} wrapper and ==== separators
    -v  / --with-choices:    Word document with questions and choices
    -nv / --no-choices:      Word document with questions only
    -fn / --from-new:        Convert new format (.txt) back to original text format
    -fc / --from-continuous: Convert ++++ separated format (continuous letters) to original

Usage:
    python quiz_processor.py input.txt                      # Generate all formats
    python quiz_processor.py -f input.txt                   # Table format only
    python quiz_processor.py -nv -v input.txt               # Questions only + with choices
    python quiz_processor.py -fn input_new_format.txt       # New format → original text
    python quiz_processor.py -fc input_continuous.txt       # Continuous format → original text
"""

import sys
import re
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt


def parse_quiz_file(filename):
    """
    Parse the input text file and extract questions and options.
    Handles both numbered (1. Question) and unnumbered questions.

    Returns:
        List of dicts: {'number', 'question', 'options': [{'letter', 'text', 'is_correct'}]}
    """
    with open(filename, 'r', encoding='utf-8') as f:
        content = f.read()

    questions = []
    current_question = None
    pending_line = None

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue

        # Option line: "a) Option text" or "a) *Option text"
        option_match = re.match(r'^([a-z])\)\s+(\*?)(.+)$', line)
        if option_match:
            if not current_question:
                if pending_line:
                    current_question = {'number': '0', 'question': pending_line, 'options': []}
                    pending_line = None
                else:
                    continue
            current_question['options'].append({
                'letter': option_match.group(1),
                'text': option_match.group(3).strip(),
                'is_correct': option_match.group(2) == '*',
            })
            continue

        # Numbered question: "1. Question text"
        question_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if question_match:
            if current_question and current_question['options']:
                questions.append(current_question)
            current_question = {
                'number': question_match.group(1),
                'question': question_match.group(2),
                'options': [],
            }
            pending_line = None
            continue

        # Unnumbered question candidate
        if current_question and current_question['options']:
            questions.append(current_question)
            current_question = None
        pending_line = line

    if current_question and current_question['options']:
        questions.append(current_question)

    for i, q in enumerate(questions, 1):
        if q['number'] == '0':
            q['number'] = str(i)

    return questions


# ---------------------------------------------------------------------------
# NEW FORMAT  →  ORIGINAL FORMAT
# ---------------------------------------------------------------------------

def parse_new_format_file(filename):
    """
    Parse a file in the hash/bracket format and return the standard question list.

    Expected structure:
        {
        Question text
        ====
        Option A
        ====
        #Correct option
        ====
        Option C
        ====

        +++++
        Next question
        ...
        }
    """
    with open(filename, 'r', encoding='utf-8') as f:
        content = f.read()

    content = content.strip()
    if content.startswith('{'):
        content = content[1:]
    if content.endswith('}'):
        content = content[:-1]
    content = content.strip()

    questions = []
    letters = 'abcdefghijklmnopqrstuvwxyz'
    q_number = 0  # only increments for valid questions

    # Normalize Windows line endings, then split on lines containing only ++++/+++++
    content = content.replace('\r\n', '\n')
    question_blocks = re.split(r'\n[ \t]*\+{4,5}[ \t]*\n', content)

    for q_idx, block in enumerate(question_blocks):
        block = block.strip()
        if not block:
            continue

        parts = re.split(r'\n[ \t]*====[ \t]*\n', block)
        parts = [p.strip() for p in parts if p.strip()]

        if not parts:
            continue

        question_text = parts[0]

        # Skip blocks that are just separators or stray markers
        if re.fullmatch(r'[+=]+', question_text) or question_text.startswith('#'):
            continue

        raw_options = parts[1:]

        options = []
        for j, raw in enumerate(raw_options):
            if j >= len(letters):
                break
            is_correct = raw.startswith('#')
            text = raw[1:].strip() if is_correct else raw.strip()
            options.append({
                'letter': letters[j],
                'text': text,
                'is_correct': is_correct,
            })

        q_number += 1
        questions.append({
            'number': str(q_number),
            'question': question_text,
            'options': options,
        })

    return questions


# ---------------------------------------------------------------------------
# CONTINUOUS FORMAT  →  ORIGINAL FORMAT
# ---------------------------------------------------------------------------

def parse_continuous_format_file(filename):
    """
    Parse a file where questions are separated by ++++ (4 pluses) and
    options use continuously incrementing letters across questions
    (a,b,c,d for Q1 then e,f,g,h for Q2, etc.).

    Option format: " e) *Correct text"  or  " e) Regular text"

    Returns the standard question list (options reset to a,b,c,d per question).
    """
    with open(filename, 'r', encoding='utf-8') as f:
        content = f.read()

    questions = []
    q_number = 0
    letters = 'abcdefghijklmnopqrstuvwxyz'

    # Robust split: handles trailing spaces, \r\n, or any surrounding whitespace
    blocks = re.split(r'\s*\+{4}\s*', content)

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        question_text = None
        raw_options = []

        for line in block.splitlines():
            line_stripped = line.strip()
            if not line_stripped:
                continue

            # Option line: "e) *text" or "e) text"
            opt_match = re.match(r'^([a-z])\)\s+(\*?)(.+)$', line_stripped)
            if opt_match:
                raw_options.append({
                    'text': opt_match.group(3).strip(),
                    'is_correct': opt_match.group(2) == '*',
                })
                continue

            # Numbered question: "1. Question text"
            num_match = re.match(r'^\d+\.\s+(.+)$', line_stripped)
            if num_match:
                question_text = num_match.group(1)
                continue

            # Unnumbered question: only set if no options seen yet
            if not raw_options:
                question_text = line_stripped

        if not question_text or not raw_options:
            continue

        q_number += 1

        # Reset option letters to a, b, c, d... regardless of original letters
        options = [
            {
                'letter': letters[i],
                'text': opt['text'],
                'is_correct': opt['is_correct'],
            }
            for i, opt in enumerate(raw_options)
        ]

        questions.append({
            'number': str(q_number),
            'question': question_text,
            'options': options,
        })

    return questions


# ---------------------------------------------------------------------------
# WRITE ORIGINAL FORMAT
# ---------------------------------------------------------------------------

def create_original_format_file(questions, output_path):
    """
    Write questions back to the plain-text original format:

        1. Question text
         a) Option 1
         b) *Correct option
         c) Option 3
         d) Option 4

        2. Next question
         ...
    """
    lines = []
    for q in questions:
        lines.append(f"{q['number']}. {q['question']}")
        for opt in q['options']:
            marker = '*' if opt['is_correct'] else ''
            lines.append(f" {opt['letter']}) {marker}{opt['text']}")
        lines.append('')

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    print(f"Created: {output_path}")


# ---------------------------------------------------------------------------
# ORIGINAL FORMAT  →  OTHER FORMATS
# ---------------------------------------------------------------------------

def create_questions_only_doc(questions, output_path):
    """Create a Word document with only questions (no options)."""
    doc = Document()
    for q in questions:
        para = doc.add_paragraph()
        run = para.add_run(f"{q['number']}. {q['question']}")
        run.font.size = Pt(12)
        para.space_after = Pt(12)
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_questions_with_choices_doc(questions, output_path):
    """Create a Word document with questions and all options (no correct-answer markers)."""
    doc = Document()
    for q in questions:
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"{q['number']}. {q['question']}")
        q_run.font.size = Pt(12)
        q_run.bold = True
        q_para.space_after = Pt(6)

        for option in q['options']:
            o_para = doc.add_paragraph()
            o_run = o_para.add_run(f"  {option['letter']}) {option['text']}")
            o_run.font.size = Pt(11)
            o_para.space_after = Pt(3)

        doc.add_paragraph().space_after = Pt(6)

    doc.save(output_path)
    print(f"Created: {output_path}")


def create_new_format_file(questions, output_path):
    """
    Create a text file in the hash/bracket format:
        {
        Question text
        ====
        Option 1
        ====
        #Correct option
        ====
        ...

        +++++

        Next question...
        }
    """
    output_lines = ['{']

    for i, q in enumerate(questions):
        output_lines.append(q['question'])
        output_lines.append('====')

        for option in q['options']:
            prefix = '#' if option['is_correct'] else ''
            output_lines.append(f"{prefix}{option['text']}")
            output_lines.append('====')

        if i < len(questions) - 1:
            output_lines.append('')
            output_lines.append('+++++')
            output_lines.append('')

    output_lines.append('}')

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(output_lines))

    print(f"Created: {output_path}")


def create_table_format_doc(questions, output_path):
    """
    Create a Word document with each question in a single-column table:
      Row 1 : Question text
      Row 2 : Correct answer
      Row 3+: Other options in original order
    """
    doc = Document()

    for q in questions:
        correct_option = next((o for o in q['options'] if o['is_correct']), None)
        other_options = [o for o in q['options'] if not o['is_correct']]

        if not correct_option and q['options']:
            correct_option = q['options'][0]
            other_options = q['options'][1:]

        if not correct_option:
            continue

        total_rows = 2 + len(other_options)
        table = doc.add_table(rows=total_rows, cols=1)
        table.style = 'Table Grid'

        q_run = table.rows[0].cells[0].paragraphs[0].add_run(q['question'])
        q_run.font.size = Pt(12)

        c_run = table.rows[1].cells[0].paragraphs[0].add_run(correct_option['text'])
        c_run.font.size = Pt(11)

        for idx, option in enumerate(other_options, start=2):
            o_run = table.rows[idx].cells[0].paragraphs[0].add_run(option['text'])
            o_run.font.size = Pt(11)

        doc.add_paragraph()

    doc.save(output_path)
    print(f"Created: {output_path}")


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='Quiz Processor - Convert multiple choice questions to various formats',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python quiz_processor.py input.txt                      # Generate all formats
  python quiz_processor.py -f input.txt                   # Table format only
  python quiz_processor.py -nv -v input.txt               # Questions only + with choices
  python quiz_processor.py -f -n -v -nv input.txt         # All formats
  python quiz_processor.py -fn input_new_format.txt       # New format → original text
  python quiz_processor.py -fc input_continuous.txt       # Continuous format → original text
        """,
    )
    parser.add_argument('input_file', help='Input file')
    parser.add_argument('-f',  '--table-format',    action='store_true', help='Word doc: table format')
    parser.add_argument('-n',  '--hash-format',     action='store_true', help='Text file: hash/bracket format')
    parser.add_argument('-v',  '--with-choices',    action='store_true', help='Word doc: questions + choices')
    parser.add_argument('-nv', '--no-choices',      action='store_true', help='Word doc: questions only')
    parser.add_argument('-fn', '--from-new',        action='store_true', help='New format → original text')
    parser.add_argument('-fc', '--from-continuous', action='store_true', help='Continuous ++++ format → original text')

    args = parser.parse_args()

    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"Error: File '{args.input_file}' not found.")
        sys.exit(1)

    output_dir = Path('quiz_output')
    output_dir.mkdir(exist_ok=True)
    base_name = input_path.stem

    # ------------------------------------------------------------------
    # MODE: new format → original
    # ------------------------------------------------------------------
    if args.from_new:
        original_output = output_dir / f"{base_name}_original.txt"
        try:
            questions = parse_new_format_file(args.input_file)
            if not questions:
                print("Warning: No questions found in the new-format file.")
                sys.exit(1)
            print(f"Parsed {len(questions)} questions from '{args.input_file}'")
            print(f"Output directory: {output_dir.absolute()}\n")
            create_original_format_file(questions, original_output)
            print(f"\nSuccess! Generated:\n  - {original_output}")
        except Exception as e:
            print(f"Error processing file: {e}")
            import traceback; traceback.print_exc()
            sys.exit(1)
        return

    # ------------------------------------------------------------------
    # MODE: continuous format → original
    # ------------------------------------------------------------------
    if args.from_continuous:
        original_output = output_dir / f"{base_name}_original.txt"
        try:
            questions = parse_continuous_format_file(args.input_file)
            if not questions:
                print("Warning: No questions found in the file.")
                sys.exit(1)
            print(f"Parsed {len(questions)} questions from '{args.input_file}'")
            print(f"Output directory: {output_dir.absolute()}\n")
            create_original_format_file(questions, original_output)
            print(f"\nSuccess! Generated:\n  - {original_output}")
        except Exception as e:
            print(f"Error processing file: {e}")
            import traceback; traceback.print_exc()
            sys.exit(1)
        return

    # ------------------------------------------------------------------
    # MODE: original → other formats
    # ------------------------------------------------------------------
    generate_all = not any([args.table_format, args.hash_format,
                            args.with_choices, args.no_choices])

    table_output          = output_dir / f"{base_name}_table_format.docx"
    hash_output           = output_dir / f"{base_name}_new_format.txt"
    with_choices_output   = output_dir / f"{base_name}_with_choices.docx"
    questions_only_output = output_dir / f"{base_name}_questions.docx"

    try:
        questions = parse_quiz_file(args.input_file)
        if not questions:
            print("Warning: No questions found in the input file.")
            sys.exit(1)

        print(f"Parsed {len(questions)} questions from '{args.input_file}'")
        print(f"Output directory: {output_dir.absolute()}\n")

        generated = []

        if generate_all or args.table_format:
            create_table_format_doc(questions, table_output)
            generated.append(table_output)

        if generate_all or args.hash_format:
            create_new_format_file(questions, hash_output)
            generated.append(hash_output)

        if generate_all or args.with_choices:
            create_questions_with_choices_doc(questions, with_choices_output)
            generated.append(with_choices_output)

        if generate_all or args.no_choices:
            create_questions_only_doc(questions, questions_only_output)
            generated.append(questions_only_output)

        print("\nSuccess! Generated documents:")
        for p in generated:
            print(f"  - {p}")

    except Exception as e:
        print(f"Error processing file: {e}")
        import traceback; traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()