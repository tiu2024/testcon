"""Output writers for all supported quiz formats.

Writers:
    write_table_doc     — Word doc with each question in a bordered table
    write_hemis_text    — Text file in hemis format (brace/====/+++++ separated)
    write_choices_doc   — Word doc with questions and all answer choices
    write_questions_doc — Word doc with numbered questions only
    write_original_text — Plain text in the standard numbered format
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt

from models import Question

_QUESTION_FONT = Pt(12)
_OPTION_FONT = Pt(11)

_HEMIS_OPTION_SEP = "===="
_HEMIS_QUESTION_SEP = "+++++"


# ---------------------------------------------------------------------------
# Word: table format
# ---------------------------------------------------------------------------

def write_table_doc(questions: list[Question], output_path: Path) -> None:
    """Each question as a single-column table: question → correct → others."""
    doc = Document()

    for question in questions:
        correct = question.correct_option
        if correct is None and question.options:
            correct = question.options[0]
        if correct is None:
            continue

        others = [o for o in question.options if o is not correct]
        table = doc.add_table(rows=2 + len(others), cols=1)
        table.style = "Table Grid"

        _set_cell(table.rows[0].cells[0], question.text, _QUESTION_FONT)
        _set_cell(table.rows[1].cells[0], correct.text, _OPTION_FONT)
        for i, option in enumerate(others, start=2):
            _set_cell(table.rows[i].cells[0], option.text, _OPTION_FONT)

        doc.add_paragraph()

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Word: questions with choices
# ---------------------------------------------------------------------------

def write_choices_doc(questions: list[Question], output_path: Path) -> None:
    """Questions followed by all answer choices (no correct-answer marker)."""
    doc = Document()

    for question in questions:
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"{question.number}. {question.text}")
        q_run.font.size = _QUESTION_FONT
        q_run.bold = True
        q_para.space_after = Pt(6)

        for option in question.options:
            o_para = doc.add_paragraph()
            o_run = o_para.add_run(f"  {option.letter}) {option.text}")
            o_run.font.size = _OPTION_FONT
            o_para.space_after = Pt(3)

        doc.add_paragraph().space_after = Pt(6)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Word: questions only
# ---------------------------------------------------------------------------

def write_questions_doc(questions: list[Question], output_path: Path) -> None:
    """Numbered questions only, no options."""
    doc = Document()

    for question in questions:
        para = doc.add_paragraph()
        run = para.add_run(f"{question.number}. {question.text}")
        run.font.size = _QUESTION_FONT
        para.space_after = Pt(12)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Text: hemis format
# ---------------------------------------------------------------------------

def write_hemis_text(questions: list[Question], output_path: Path) -> None:
    """Brace-wrapped text with ==== and +++++ separators; # marks correct."""
    lines: list[str] = ["{"]

    for i, question in enumerate(questions):
        lines.append(question.text)
        lines.append(_HEMIS_OPTION_SEP)

        for option in question.options:
            prefix = "#" if option.is_correct else ""
            lines.append(f"{prefix}{option.text}")
            lines.append(_HEMIS_OPTION_SEP)

        if i < len(questions) - 1:
            lines.extend(["", _HEMIS_QUESTION_SEP, ""])

    lines.append("}")
    output_path.write_text("\n".join(lines), encoding="utf-8")


# ---------------------------------------------------------------------------
# Text: original format
# ---------------------------------------------------------------------------

def write_original_text(questions: list[Question], output_path: Path) -> None:
    """Standard numbered format with lettered options (* marks correct)."""
    lines: list[str] = []

    for question in questions:
        lines.append(f"{question.number}. {question.text}")
        for option in question.options:
            marker = "*" if option.is_correct else ""
            lines.append(f" {option.letter}) {marker}{option.text}")
        lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell(cell, text: str, font_size: Pt) -> None:
    run = cell.paragraphs[0].add_run(text)
    run.font.size = font_size