#!/usr/bin/env python3
"""
Quiz Processor - Converts multiple choice questions from text to Word documents and new format.

Generates output files based on flags:
    -f / --table-format: Word document with questions in table format
    -n / --hash-format: Text file with {} wrapper and ==== separators
    -v / --with-choices: Word document with questions and choices
    -nv / --no-choices: Word document with questions only

Usage:
    python quiz_processor.py input.txt              # Generate all formats
    python quiz_processor.py -f input.txt           # Table format only
    python quiz_processor.py -nv -v input.txt       # Questions only + with choices
"""
 
import sys
import re
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT


def parse_quiz_file(filename):
    """
    Parse the input text file and extract questions and options.
    Handles both numbered (1. Question) and unnumbered questions.
    
    Returns:
        List of dictionaries with structure:
        {
            'number': question number,
            'question': question text,
            'options': [{'letter': 'a', 'text': 'option text', 'is_correct': False}, ...]
        }
    """
    with open(filename, 'r', encoding='utf-8') as f:
        content = f.read()
    
    questions = []
    current_question = None
    pending_line = None  # Store potential unnumbered question
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
        
        # Check if this is an option line: "a) Option text" or "a) *Option text"
        option_match = re.match(r'^([a-z])\)\s+(\*?)(.+)$', line)
        
        if option_match:
            # If we encounter an option but don't have a current question
            if not current_question:
                # Check if we have a pending line (unnumbered question)
                if pending_line:
                    current_question = {
                        'number': '0',  # Will be auto-numbered later
                        'question': pending_line,
                        'options': []
                    }
                    pending_line = None
                else:
                    # Skip orphaned options
                    continue
            
            # Add the option to current question
            letter = option_match.group(1)
            is_correct = option_match.group(2) == '*'
            text = option_match.group(3).strip()
            
            current_question['options'].append({
                'letter': letter,
                'text': text,
                'is_correct': is_correct
            })
            continue
        
        # Try to match numbered question: "1. Question text"
        question_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if question_match:
            # Save the previous question if it exists
            if current_question and current_question['options']:
                questions.append(current_question)
            
            current_question = {
                'number': question_match.group(1),
                'question': question_match.group(2),
                'options': []
            }
            pending_line = None
            continue
        
        # If we reach here, this is neither a numbered question nor an option
        # It might be an unnumbered question
        
        # Save the previous question if it has options
        if current_question and current_question['options']:
            questions.append(current_question)
            current_question = None
        
        # Store this line as a potential question
        pending_line = line
    
    # Don't forget the last question
    if current_question and current_question['options']:
        questions.append(current_question)
    
    # Auto-number questions that don't have numbers
    for i, q in enumerate(questions, 1):
        if q['number'] == '0':
            q['number'] = str(i)
    
    return questions


def create_questions_only_doc(questions, output_path):
    """Create a Word document with only questions (no options)."""
    doc = Document()
    
    for q in questions:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"{q['number']}. {q['question']}")
        run.font.size = Pt(12)
        
        # Add spacing after each question
        paragraph.space_after = Pt(12)
    
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_questions_with_choices_doc(questions, output_path):
    """Create a Word document with questions and all options (without correct answer markers)."""
    doc = Document()
    
    for q in questions:
        # Add question
        question_para = doc.add_paragraph()
        question_run = question_para.add_run(f"{q['number']}. {q['question']}")
        question_run.font.size = Pt(12)
        question_run.bold = True
        question_para.space_after = Pt(6)
        
        # Add options (no bullet points)
        for option in q['options']:
            option_para = doc.add_paragraph()
            option_run = option_para.add_run(f"  {option['letter']}) {option['text']}")
            option_run.font.size = Pt(11)
            option_para.space_after = Pt(3)
        
        # Add spacing after each question block
        doc.add_paragraph().space_after = Pt(6)
    
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_new_format_file(questions, output_path):
    """
    Create a text file with questions in the new format:
    {
    Question text
    ====
    Option 1
    ====
    #Correct option
    ====
    Option 3
    ====
    
    +++++
    Next question
    ...
    }
    """
    output_lines = ['{']
    
    for i, q in enumerate(questions):
        # Add question text
        output_lines.append(q['question'])
        output_lines.append('====')
        
        # Add options
        for option in q['options']:
            # Prefix with # if it's the correct answer
            prefix = '#' if option['is_correct'] else ''
            output_lines.append(f"{prefix}{option['text']}")
            output_lines.append('====')
        
        # Add separator between questions (but not after the last one)
        if i < len(questions) - 1:
            output_lines.append('')
            output_lines.append('+++++')
        
        # Add empty line after separator for readability
        if i < len(questions) - 1:
            output_lines.append('')
    
    output_lines.append('}')
    
    # Write to file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(output_lines))
    
    print(f"Created: {output_path}")


def create_table_format_doc(questions, output_path):
    """
    Create a Word document with each question in a table format.
    
    Table structure:
    - Row 1: Question text
    - Row 2: Correct answer (just text, no label) or first option if no correct answer
    - Row 3+: Other options in original order (just text, no labels)
    """
    doc = Document()
    
    for q in questions:
        # Find the correct answer and other options
        correct_option = None
        other_options = []
        
        for option in q['options']:
            if option['is_correct']:
                correct_option = option
            else:
                other_options.append(option)
        
        # If no correct answer is marked, use first option as placeholder
        if not correct_option and q['options']:
            correct_option = q['options'][0]
            other_options = q['options'][1:]
        
        # Skip questions with no options
        if not correct_option:
            continue
        
        # Calculate total rows: 1 (question) + 1 (correct) + len(other_options)
        total_rows = 2 + len(other_options)
        
        # Create table with single column (simple, no styling)
        table = doc.add_table(rows=total_rows, cols=1)
        table.style = 'Table Grid'
        
        # Row 1: Question
        question_cell = table.rows[0].cells[0]
        question_para = question_cell.paragraphs[0]
        question_run = question_para.add_run(q['question'])
        question_run.font.size = Pt(12)
        
        # Row 2: Correct answer
        correct_cell = table.rows[1].cells[0]
        correct_para = correct_cell.paragraphs[0]
        correct_run = correct_para.add_run(correct_option['text'])
        correct_run.font.size = Pt(11)
        
        # Remaining rows: Other options in original order
        row_idx = 2
        for option in other_options:
            option_cell = table.rows[row_idx].cells[0]
            option_para = option_cell.paragraphs[0]
            option_run = option_para.add_run(option['text'])
            option_run.font.size = Pt(11)
            row_idx += 1
        
        # Add spacing after each table
        doc.add_paragraph()
    
    doc.save(output_path)
    print(f"Created: {output_path}")


def main():
    """Main function to process the quiz file."""
    # Parse command-line arguments
    parser = argparse.ArgumentParser(
        description='Quiz Processor - Convert multiple choice questions to various formats',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python quiz_processor.py input.txt              # Generate all formats
  python quiz_processor.py -f input.txt           # Table format only
  python quiz_processor.py -nv -v input.txt       # Questions only + with choices
  python quiz_processor.py -f -n -v -nv input.txt # All formats
        """
    )
    parser.add_argument('input_file', help='Input text file with questions')
    parser.add_argument('-f', '--table-format', action='store_true',
                       help='Generate Word document with table format')
    parser.add_argument('-n', '--hash-format', action='store_true',
                       help='Generate text file with hash/bracket format')
    parser.add_argument('-v', '--with-choices', action='store_true',
                       help='Generate Word document with questions and choices')
    parser.add_argument('-nv', '--no-choices', action='store_true',
                       help='Generate Word document with questions only')
    
    args = parser.parse_args()
    
    # Check if input file exists
    if not Path(args.input_file).exists():
        print(f"Error: File '{args.input_file}' not found.")
        sys.exit(1)
    
    # Determine which formats to generate
    generate_all = not any([args.table_format, args.hash_format, 
                           args.with_choices, args.no_choices])
    
    # Create quiz_output directory
    output_dir = Path('quiz_output')
    output_dir.mkdir(exist_ok=True)
    
    # Generate output filenames based on input filename
    input_path = Path(args.input_file)
    base_name = input_path.stem  # filename without extension
    
    table_output = output_dir / f"{base_name}_table_format.docx"
    hash_output = output_dir / f"{base_name}_new_format.txt"
    with_choices_output = output_dir / f"{base_name}_with_choices.docx"
    questions_only_output = output_dir / f"{base_name}_questions.docx"
    
    # Parse the quiz file
    try:
        questions = parse_quiz_file(args.input_file)
        
        if not questions:
            print("Warning: No questions found in the input file.")
            sys.exit(1)
        
        print(f"Parsed {len(questions)} questions from '{args.input_file}'")
        print(f"Output directory: {output_dir.absolute()}\n")
        
        generated_files = []
        
        # Generate requested formats
        if generate_all or args.table_format:
            create_table_format_doc(questions, table_output)
            generated_files.append(table_output)
        
        if generate_all or args.hash_format:
            create_new_format_file(questions, hash_output)
            generated_files.append(hash_output)
        
        if generate_all or args.with_choices:
            create_questions_with_choices_doc(questions, with_choices_output)
            generated_files.append(with_choices_output)
        
        if generate_all or args.no_choices:
            create_questions_only_doc(questions, questions_only_output)
            generated_files.append(questions_only_output)
        
        print("\nSuccess! Generated documents:")
        for file_path in generated_files:
            print(f"  - {file_path}")
        
    except Exception as e:
        print(f"Error processing file: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
